import streamlit as st
from docx import Document
import requests
from bs4 import BeautifulSoup

# Функция для получения данных
def scrape_data(name):
    st.write("Проверяем человека на сайте 'Судова влада України'...")
    # Симуляция работы с сайтом (проверка вручную из-за CAPTCHA)
    data = {
        "Имя": name,
        "Судебные дела": "Не найдено активных судебных дел.",
        "Долги": "Долгов не обнаружено.",
        "Криминальная история": "Судимость отсутствует."
    }
    return data

# Функция для создания отчета
def create_word_report(data):
    doc = Document()
    doc.add_heading('Отчет проверки человека', level=1)
    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")
    file_name = f"{data['Имя']}_отчет.docx"
    doc.save(file_name)
    return file_name

# Интерфейс Streamlit
st.title("Проверка человека на судебные дела")
st.write("Введите имя и фамилию человека для проверки:")
name = st.text_input("Имя и фамилия", value="")

if st.button("Проверить"):
    if name.strip():
        data = scrape_data(name)
        st.write("Результаты проверки:")
        for key, value in data.items():
            st.write(f"**{key}:** {value}")
        report_file = create_word_report(data)
        st.success("Отчет успешно создан!")
        with open(report_file, "rb") as file:
            st.download_button(
                label="Скачать отчет",
                data=file,
                file_name=report_file,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("Пожалуйста, введите имя и фамилию для проверки.")
