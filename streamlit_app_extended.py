
import streamlit as st
from docx import Document
import requests
from bs4 import BeautifulSoup

# Function to scrape public data
def scrape_data(name, id_number=None, birth_date=None, city=None):
    # Simulate scraping data with additional fields
    data = {
        "Name": name,
        "ID Number": id_number or "Not provided",
        "Birth Date": birth_date or "Not provided",
        "City": city or "Not provided",
        "Court Cases": "No active court cases found.",
        "Debts": "No outstanding debts found.",
        "Criminal Record": "No criminal records found."
    }
    return data

# Function to create Word document
def create_word_report(data):
    doc = Document()
    doc.add_heading("Перевірка особи: Звіт", level=1)
    doc.add_paragraph(f"Ім'я: {data['Name']}")
    doc.add_paragraph(f"Ідентифікаційний код: {data['ID Number']}")
    doc.add_paragraph(f"Дата народження: {data['Birth Date']}")
    doc.add_paragraph(f"Місто проживання: {data['City']}")
    doc.add_paragraph(f"Судові справи: {data['Court Cases']}")
    doc.add_paragraph(f"Борги: {data['Debts']}")
    doc.add_paragraph(f"Кримінальні записи: {data['Criminal Record']}")
    doc.save(f"{data['Name']}_звіт.docx")

# Streamlit app
st.title("Перевірка особи на судові справи")
st.write("Цей застосунок допомагає перевірити особу на наявність судових справ, боргів та інших даних.")

name = st.text_input("Введіть ім'я та прізвище особи для перевірки:")
id_number = st.text_input("Введіть ідентифікаційний код (за бажанням):")
birth_date = st.text_input("Введіть дату народження (формат: РРРР-ММ-ДД, за бажанням):")
city = st.text_input("Введіть місто проживання (за бажанням):")

if st.button("Перевірити"):
    if not name:
        st.error("Будь ласка, введіть ім'я та прізвище для перевірки!")
    else:
        st.write(f"Перевіряємо особу: {name}...")
        data = scrape_data(name, id_number, birth_date, city)
        st.write("Результати перевірки:")
        st.write(f"Ім'я: {data['Name']}")
        st.write(f"Ідентифікаційний код: {data['ID Number']}")
        st.write(f"Дата народження: {data['Birth Date']}")
        st.write(f"Місто проживання: {data['City']}")
        st.write(f"Судові справи: {data['Court Cases']}")
        st.write(f"Борги: {data['Debts']}")
        st.write(f"Кримінальні записи: {data['Criminal Record']}")
        create_word_report(data)
        st.success("Звіт створено! Ви можете завантажити його.")
