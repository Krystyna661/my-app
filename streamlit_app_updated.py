
import streamlit as st
from docx import Document
import requests
from bs4 import BeautifulSoup

def scrape_data(name, id_number=None, city=None):
    # Simulated scraping logic
    data = {
        "Name": name,
        "ID Number": id_number or "Not Provided",
        "City": city or "Not Provided",
        "Court Cases": "No active court cases found.",
        "Debts": "No outstanding debts found.",
        "Criminal Record": "No criminal records found."
    }
    return data

def create_word_report(data):
    doc = Document()
    doc.add_heading('Report', level=1)
    doc.add_paragraph(f"Name: {data['Name']}")
    doc.add_paragraph(f"ID Number: {data['ID Number']}")
    doc.add_paragraph(f"City: {data['City']}")
    doc.add_paragraph(f"Court Cases: {data['Court Cases']}")
    doc.add_paragraph(f"Debts: {data['Debts']}")
    doc.add_paragraph(f"Criminal Record: {data['Criminal Record']}")
    doc.save(f"{data['Name']}_Report.docx")

st.title("Person Background Checker")
st.write("This application helps you check a person's background, including court cases, debts, and criminal records.")

name = st.text_input("Enter the person's name:")
id_number = st.text_input("Enter the person's ID number (optional):")
city = st.text_input("Enter the person's city (optional):")

if st.button("Check"):
    if name:
        st.write(f"Checking background for: {name}")
        data = scrape_data(name, id_number, city)
        st.write("Results:")
        st.write(data)
        create_word_report(data)
        st.success(f"Report for {name} generated successfully!")
    else:
        st.error("Please enter a name to proceed.")
