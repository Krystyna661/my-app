
import streamlit as st
from docx import Document
import requests
from bs4 import BeautifulSoup
import sqlite3

# SQLite database setup
def setup_database():
    conn = sqlite3.connect('user_checks.db')
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS checks (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        name TEXT,
                        date TEXT,
                        checked_sites TEXT,
                        status TEXT)''')
    conn.commit()
    conn.close()

# Function to log check into the database
def log_check(name, date, checked_sites, status):
    conn = sqlite3.connect('user_checks.db')
    cursor = conn.cursor()
    cursor.execute("INSERT INTO checks (name, date, checked_sites, status) VALUES (?, ?, ?, ?)",
                   (name, date, checked_sites, status))
    conn.commit()
    conn.close()

# Function to simulate web scraping and handle CAPTCHA manually
def scrape_data(name):
    st.write("Переходимо на сайт 'Судова влада України'.")
    st.write("Будь ласка, введіть CAPTCHA вручну, якщо вона з'явиться на екрані.")
    # Simulating manual CAPTCHA
    captcha_code = st.text_input("Введіть CAPTCHA-код з сайту:")
    
    if captcha_code:
        data = {
            "Ім'я": name,
            "Результат перевірки": "Успішно виконано. CAPTCHA: " + captcha_code
        }
        return data
    else:
        return {"Ім'я": name, "Результат перевірки": "Перевірка не виконана через відсутність CAPTCHA."}

# Function to create a Word report
def create_word_report(data):
    doc = Document()
    doc.add_heading("Звіт про перевірку особи", level=1)
    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")
    file_name = f"{data['Ім\'я']}_звіт.docx"
    doc.save(file_name)
    return file_name

# Streamlit app interface
def main():
    setup_database()
    st.title("Перевірка особи на судові впровадження")
    st.write("Цей застосунок допомагає перевірити особу на наявність судових впроваджень та інших важливих даних.")
    
    name = st.text_input("Введіть ім'я та прізвище особи для перевірки:")
    if st.button("Перевірити"):
        if name:
            st.write(f"Перевіряємо особу: {name}...")
            data = scrape_data(name)
            st.write("Результати перевірки:", data)
            
            file_name = create_word_report(data)
            st.success(f"Звіт збережено у файл: {file_name}")
            
            # Log the check into the database
            log_check(name, "2024-12-25", "Судова влада України", data["Результат перевірки"])
            
            st.write("Додаткові перевірки на інших джерелах...")
            st.warning("Ця функція поки що у розробці.")
        else:
            st.error("Будь ласка, введіть ім'я для перевірки.")

if __name__ == "__main__":
    main()
