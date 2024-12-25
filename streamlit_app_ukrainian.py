
import streamlit as st
from docx import Document
import requests
from bs4 import BeautifulSoup

def scrape_data(name):
    # Імітація збору даних з кількох джерел
    data = {
        "Ім'я": name,
        "Судові справи": "Не знайдено активних судових справ.",
        "Борги": "Боргів не виявлено.",
        "Кримінальні записи": "Кримінальних записів не знайдено.",
        "Перевірені джерела": [
            "Судова влада України",
            "Єдиний реєстр боржників",
            "Інші публічні бази даних"
        ]
    }
    return data

def create_word_report(data):
    doc = Document()
    doc.add_heading('Звіт про перевірку особи', level=1)
    doc.add_paragraph(f"Ім'я: {data['Ім'я']}")
    doc.add_paragraph(f"Судові справи: {data['Судові справи']}")
    doc.add_paragraph(f"Борги: {data['Борги']}")
    doc.add_paragraph(f"Кримінальні записи: {data['Кримінальні записи']}")
    doc.add_paragraph("Перевірені джерела:")
    for source in data['Перевірені джерела']:
        doc.add_paragraph(f"- {source}")
    doc.save(f"{data['Ім'я']}_звіт.docx")

st.title("Перевірка кандидата")
st.write("Цей застосунок допомагає аналізувати публічні дані для оцінки ризиків під час найму.")

name = st.text_input("Введіть ім'я та прізвище особи для перевірки:")

if st.button("Перевірити"):
    if name:
        st.write(f"Перевіряємо особу: {name}...")
        data = scrape_data(name)
        st.write("Результати перевірки:")
        for key, value in data.items():
            if key != "Перевірені джерела":
                st.write(f"{key}: {value}")
        st.write("Створення звіту...")
        create_word_report(data)
        st.success(f"Звіт збережено: {data['Ім'я']}_звіт.docx")
    else:
        st.error("Будь ласка, введіть ім'я для перевірки!")
